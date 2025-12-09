import tkinter as tk
from tkinter import messagebox, ttk
import mysql.connector
from mysql.connector import Error

from docx import Document
from docx.shared import Inches
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
import os


class MySQLLoginApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Подключение к MySQL")
        self.root.geometry("400x300")
        self.root.resizable(False, False)

        tk.Label(root, text="Хост (сервер):").pack(pady=(10, 0))
        self.host_entry = tk.Entry(root, width=50)
        self.host_entry.insert(0, "localhost")
        self.host_entry.pack(pady=5)

        tk.Label(root, text="Имя базы данных:").pack()
        self.db_entry = tk.Entry(root, width=50)
        self.db_entry.insert(0, "test")
        self.db_entry.pack(pady=5)

        tk.Label(root, text="Пользователь:").pack()
        self.user_entry = tk.Entry(root, width=50)
        self.user_entry.insert(0, "root")
        self.user_entry.pack(pady=5)

        tk.Label(root, text="Пароль:").pack()
        self.pass_entry = tk.Entry(root, width=50, show="*")
        self.pass_entry.pack(pady=5)

        self.connect_btn = tk.Button(
            root,
            text="Подключиться",
            command=self.connect_to_db,
            bg="#4CAF50",
            fg="white",
            font=("Arial", 10, "bold")
        )
        self.connect_btn.pack(pady=20)

        self.status_label = tk.Label(root, text="", fg="red", font=("Arial", 9))
        self.status_label.pack()

        self.connection = None

    def connect_to_db(self):
        host = self.host_entry.get().strip()
        database = self.db_entry.get().strip()
        username = self.user_entry.get().strip()
        password = self.pass_entry.get().strip()

        if not all([host, database, username]):
            messagebox.showwarning("Ошибка", "Хост, база данных и пользователь обязательны!")
            return

        try:
            self.connection = mysql.connector.connect(
                host=host,
                database=database,
                user=username,
                password=password
            )

            if self.connection.is_connected():
                messagebox.showinfo("Успех", f"Успешное подключение к базе '{database}'!")
                self.status_label.config(text="Подключено", fg="green")
                self.create_data_window()

        except Error as e:
            messagebox.showerror("Ошибка подключения", f"Не удалось подключиться:\n{e}")
            self.status_label.config(text="Ошибка", fg="red")
            self.connection = None

    def create_data_window(self):
        data_window = tk.Toplevel(self.root)
        data_window.title("Меню")
        data_window.geometry("800x600")
        data_window.resizable(True, True)

        # Фрейм слева — форма
        form_frame = tk.Frame(data_window, width=250, padx=10, pady=10, relief="groove", bd=1)
        form_frame.pack(side=tk.LEFT, fill=tk.Y, padx=5, pady=5)

        tk.Label(form_frame, text="Управление данными", font=("Arial", 12, "bold")).pack(pady=(0, 10))

        # Поля ввода
        tk.Label(form_frame, text="ID:").pack(anchor="w")
        self.id_entry = tk.Entry(form_frame, width=25)
        self.id_entry.pack(pady=2, fill=tk.X)

        tk.Label(form_frame, text="Наименование:").pack(anchor="w", pady=(10, 0))
        self.name_entry = tk.Entry(form_frame, width=25)
        self.name_entry.pack(pady=2, fill=tk.X)

        tk.Label(form_frame, text="Категория:").pack(anchor="w", pady=(10, 0))
        self.cat_entry = tk.Entry(form_frame, width=25)
        self.cat_entry.pack(pady=2, fill=tk.X)

        btn_frame = tk.Frame(form_frame)
        btn_frame.pack(pady=15)

        tk.Button(btn_frame, text="Добавить", command=self.add_product, bg="#4CAF50", fg="gray", width=12).pack(pady=2)
        tk.Button(btn_frame, text="Изменить", command=self.update_product, bg="#FF9800", fg="gray", width=12).pack(pady=2)
        tk.Button(btn_frame, text="Удалить", command=self.delete_product, bg="#F44336", fg="gray", width=12).pack(pady=2)

        tk.Button(btn_frame, text="В Word", command=self.export_to_word, bg="#2196F3", fg="gray", width=12).pack(pady=2)
        tk.Button(btn_frame, text="В PDF", command=self.export_to_pdf, bg="#9C27B0", fg="gray", width=12).pack(pady=2)

        # Фрейм справа — таблица

        table_frame = tk.Frame(data_window)
        table_frame.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True, padx=5, pady=5)

        self.tree_view = ttk.Treeview(table_frame, show="headings")
        self.tree_view.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        vsb = ttk.Scrollbar(table_frame, orient="vertical", command=self.tree_view.yview)
        vsb.pack(side="right", fill="y")
        self.tree_view.configure(yscrollcommand=vsb.set)

        hsb = ttk.Scrollbar(data_window, orient="horizontal", command=self.tree_view.xview)
        hsb.pack(side="bottom", fill="x")
        self.tree_view.configure(xscrollcommand=hsb.set)

        # Загрузка данных
        self.load_product_data(self.tree_view)

        # Обработчик клика по строке
        self.tree_view.bind("<ButtonRelease-1>", self.on_row_select)

    def load_product_data(self, tree):
        if not self.connection or not self.connection.is_connected():
            return

        try:
            cursor = self.connection.cursor()
            cursor.execute("SELECT * FROM product")
            rows = cursor.fetchall()
            columns = [col[0] for col in cursor.description]

            for item in tree.get_children():
                tree.delete(item)

            tree["columns"] = columns
            for col in columns:
                tree.heading(col, text=col)
                tree.column(col, width=100, anchor=tk.CENTER)

            for row in rows:
                tree.insert("", "end", values=row)

        except Error as e:
            messagebox.showerror("Ошибка", f"Не удалось загрузить данные:\n{e}")
        finally:
            cursor.close()

    def on_row_select(self, event):
        selected = self.tree_view.selection()
        if selected:
            values = self.tree_view.item(selected[0], "values")
            self.id_entry.delete(0, tk.END)
            self.id_entry.insert(0, values[0])
            self.name_entry.delete(0, tk.END)
            self.name_entry.insert(0, values[1])
            self.cat_entry.delete(0, tk.END)
            self.cat_entry.insert(0, values[2])

    def add_product(self):
        name = self.name_entry.get().strip()
        cat_id = self.cat_entry.get().strip()

        if not name or not cat_id:
            messagebox.showwarning("Ошибка", "Заполните Наименование и Категорию!")
            return

        try:
            cursor = self.connection.cursor()
            cursor.execute("INSERT INTO product (name, cat_ID) VALUES (%s, %s)", (name, cat_id))
            self.connection.commit()
            messagebox.showinfo("Успех", "Товар добавлен!")
            self.load_product_data(self.tree_view)
            self.clear_entries()
        except Error as e:
            messagebox.showerror("Ошибка", f"Не удалось добавить товар:\n{e}")
        finally:
            cursor.close()

    def update_product(self):
        try:
            prod_id = int(self.id_entry.get())
            name = self.name_entry.get().strip()
            cat_id = self.cat_entry.get().strip()

            if not name or not cat_id:
                messagebox.showwarning("Ошибка", "Заполните Наименование и Категорию!")
                return

            cursor = self.connection.cursor()
            cursor.execute("UPDATE product SET name = %s, cat_ID = %s WHERE ID = %s", (name, cat_id, prod_id))
            self.connection.commit()

            if cursor.rowcount == 0:
                messagebox.showwarning("Ошибка", "Товар с таким ID не найден!")
            else:
                messagebox.showinfo("Успех", "Товар обновлён!")
                self.load_product_data(self.tree_view)
                self.clear_entries()

        except ValueError:
            messagebox.showerror("Ошибка", "ID должен быть целым числом!")
        except Error as e:
            messagebox.showerror("Ошибка", f"Не удалось обновить товар:\n{e}")
        finally:
            cursor.close()

    def delete_product(self):
        try:
            prod_id = int(self.id_entry.get())

            cursor = self.connection.cursor()
            cursor.execute("DELETE FROM product WHERE ID = %s", (prod_id,))
            self.connection.commit()

            if cursor.rowcount == 0:
                messagebox.showwarning("Ошибка", "Товар с таким ID не найден!")
            else:
                messagebox.showinfo("Успех", "Товар удалён!")
                self.load_product_data(self.tree_view)
                self.clear_entries()

        except ValueError:
            messagebox.showerror("Ошибка", "ID должен быть целым числом!")
        except Error as e:
            messagebox.showerror("Ошибка", f"Не удалось удалить товар:\n{e}")
        finally:
            cursor.close()

    def clear_entries(self):
        self.id_entry.delete(0, tk.END)
        self.name_entry.delete(0, tk.END)
        self.cat_entry.delete(0, tk.END)

    def export_to_word(self):
        if not self.tree_view.get_children():
            messagebox.showwarning("Предупреждение", "Нет данных для экспорта!")
            return

        try:
            doc = Document()
            doc.add_heading('table product', 0)

            # колонки и строки
            columns = list(self.tree_view["columns"])
            table = doc.add_table(rows=1, cols=len(columns))
            table.style = 'Table Grid'

            # Заголовки
            hdr_cells = table.rows[0].cells
            for i, col in enumerate(columns):
                hdr_cells[i].text = str(col)

            # Данные
            for row_id in self.tree_view.get_children():
                row = self.tree_view.item(row_id)["values"]
                row_cells = table.add_row().cells
                for i, val in enumerate(row):
                    row_cells[i].text = str(val)

            filepath = "exported_data.docx"
            doc.save(filepath)
            messagebox.showinfo("Успех", f"Данные экспортированы в Word:\n{os.path.abspath(filepath)}")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось экспортировать в Word:\n{e}")

    def export_to_pdf(self):
        if not self.tree_view.get_children():
            messagebox.showwarning("Предупреждение", "Нет данных для экспорта!")
            return

        try:
            filepath = "exported_data.pdf"
            doc = SimpleDocTemplate(filepath, pagesize=letter)
            elements = []

            styles = getSampleStyleSheet()
            title = Paragraph("table product", styles['Title'])
            elements.append(title)
            elements.append(Spacer(1, 12))

            # Заголовки
            columns = list(self.tree_view["columns"])
            data = [columns]

            # Данные
            for row_id in self.tree_view.get_children():
                row = self.tree_view.item(row_id)["values"]
                data.append([str(val) for val in row])

            # Создание таблицы
            table = Table(data)
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ]))

            elements.append(table)
            doc.build(elements)
            messagebox.showinfo("Успех", f"Данные экспортированы в PDF:\n{os.path.abspath(filepath)}")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось экспортировать в PDF:\n{e}")


if __name__ == "__main__":
    root = tk.Tk()
    app = MySQLLoginApp(root)
    root.mainloop()